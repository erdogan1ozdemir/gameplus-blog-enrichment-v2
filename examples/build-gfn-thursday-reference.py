# -*- coding: utf-8 -*-
"""GFN Thursday 2 Temmuz 2026 — ham docx -> v10.1 enriched. Yazar metni KORUNUR; checkpoint'li."""
import sys, os, re, zipfile
import xml.etree.ElementTree as ET
from docx import Document
sys.path.insert(0, "<skill>/scripts")
from gameplus_blog_components import *

OUT = "/Users/Erdo/Desktop/Claude Projects/Dispatch"
DOCX = "/Users/Erdo/Downloads/GFN Thursday - GeForce NOW’da Bu Hafta (2 Temmuz 2026).docx"
TITLE = "Temmuz Boyu Oyun Keyfi: 12 Yeni Oyun GeForce NOW'da!"
SLUG = "gfn-thursday-geforce-now-da-bu-hafta-2-temmuz-2026"

NS = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
      'r':'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
def Wt(t): return '{%s}%s' % (NS['w'], t)
def Rt(t): return '{%s}%s' % (NS['r'], t)
z = zipfile.ZipFile(DOCX)
rels = {rel.get('Id'): rel.get('Target') for rel in ET.fromstring(z.read('word/_rels/document.xml.rels'))}
def esc(t): return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
def run_html(r):
    txt = ''.join(n.text or '' for n in r.iter(Wt('t')))
    if not txt: return ''
    rpr = r.find(Wt('rPr'))
    b = rpr is not None and rpr.find(Wt('b')) is not None
    h = esc(txt)
    return f'<strong>{h}</strong>' if b else h
def para_parts(p):
    """(html, links[]) — hyperlink'ler <a> olarak korunur."""
    parts, links = [], []
    for ch in p._p:
        tag = ch.tag.split('}')[-1]
        if tag == 'r': parts.append(run_html(ch))
        elif tag == 'hyperlink':
            url = rels.get(ch.get(Rt('id')), '')
            inner = ''.join(run_html(r) for r in ch.findall(Wt('r')))
            if url:
                links.append(url); parts.append(f'<a href="{url}">{inner}</a>')
            else: parts.append(inner)
    return ''.join(parts), links

d = Document(DOCX)
paras = []
for p in d.paragraphs:
    if not p.text.strip(): continue
    html, links = para_parts(p)
    paras.append({'style': p.style.name, 'text': p.text.strip(), 'html': html.strip(), 'links': links})

# ---- Tür atamaları (kaynak: Figma mock tabloları — uydurma değil, tasarımın verisi) ----
GENRES = {
 'Monopoly': ('Strateji','Aile'), 'Meccha Chameleon': ('Aksiyon','Bağımsız'),
 'Assassin': ('Aksiyon','Macera'), 'Denshattack': ('Bağımsız','Aksiyon'),
 'The Mound': ('Korku',), 'Heave Ho 2': ('Parti','Bağımsız'),
 'Fogpiercer': ('Strateji','Roguelike'), 'ZeroSpace': ('Gerçek Zamanlı Strateji',),
 'The Planet Crafter': ('Hayatta Kalma',), 'Carnival Hunt': ('Korku','Co-op'),
 'The Ranchers': ('Simülasyon',), 'Corsair Cove': ('Macera','Co-op'),
 'Deer & Boy': ('Bulmaca','Bağımsız'), 'DOOM Eternal': ('FPS','Aksiyon'),
 'Embers of the Uncrowned': ('RPG',), 'EMPULSE': ('Aksiyon',),
 'The Elder Scrolls Online': ('MMORPG',), 'NBA THE RUN': ('Spor',),
 'SAND': ('Aksiyon','MMO'), 'Voidling Bound': ('Roguelike',),
 'Witchspire': ('Aksiyon','Roguelike'), 'World of Tanks': ('Aksiyon','MMO'),
}
def genres_for(name):
    for k, v in GENRES.items():
        if name.startswith(k): return v
    return ()

# ---- Stüdyo · Yıl (kaynak: Steam API appdetails / doğrulanmış: WoT=Wargaming basını, ACBF=Ubisoft, ESO=2014 resmi çıkış) ----
STUDIOS = {
 'Monopoly': 'Behaviour Interactive · 2026', 'Meccha Chameleon': 'lemorion_1224 · 2026',
 'Assassin': 'Ubisoft Singapore · 2026', 'Denshattack': 'Undercoders · 2026',
 'The Mound': 'ACE Team · 2026', 'Heave Ho 2': 'Le Cartel Studio · 2026',
 'Fogpiercer': 'Mad Cookies Studio · 2026', 'ZeroSpace': 'Starlance Studios · 2026',
 'The Planet Crafter': 'Miju Games · 2024', 'Carnival Hunt': 'Crytivo · 2026',
 'The Ranchers': 'The Ranchers · 2026', 'Corsair Cove': 'Limbic Entertainment · 2026',
 'Deer & Boy': 'Lifeline Games · 2026', 'DOOM Eternal': 'id Software · 2020',
 'Embers of the Uncrowned': 'NEXON · 2026', 'EMPULSE': '1047 Games · 2026',
 'The Elder Scrolls Online': 'ZeniMax Online Studios · 2014', 'NBA THE RUN': 'Play by Play Studios · 2026',
 'SAND': 'Hologryph · 2026', 'Voidling Bound': 'Hatchery Games · 2026',
 'Witchspire': 'Envar Games · 2026', 'World of Tanks': 'Wargaming · 2026',
}
def studio_for(name):
    for k, v in STUDIOS.items():
        if name.startswith(k): return v
    return None

# Doc'taki hatalı link düzeltmesi (rapor edildi): Embers satırındaki Steam linki ACBF sayfasına gidiyordu.
LINK_FIX = {'https://store.steampowered.com/app/3751950?utm_source=nvidia&utm_campaign=geforce_now':
            'https://store.steampowered.com/app/3767660/Embers_of_the_Uncrowned/'}

GAME_RE = re.compile(r'^(?P<name>[^()]{2,70}?)\s*\((?P<meta>[^)]{2,120})\)$')
def is_game_line(pp):
    return pp['style'] == 'normal' and GAME_RE.match(pp['text']) and not pp['text'].startswith('GFN Thursday')
def platform_of(url):
    if 'steampowered.com' in url: return 'Steam'
    if 'xbox.com' in url: return 'Xbox'
    if 'epicgames.com' in url: return 'Epic Games Store'
    if 'ubisoft.com' in url: return 'Ubisoft'
    if 'wotheat.com' in url or 'wargaming' in url: return 'Wargaming'
    return None

def linkify_meta(meta_html, links):
    """Platform/çıkış hücresi: doc'ta o satır için link varsa ilgili platform sözcüğünü o mağaza
    sayfasına linkle (↗ ikonlu, altı çizgisiz, renk inherit). Kural 8."""
    out = meta_html
    for url in links:
        plat = platform_of(url)
        if not plat: continue
        # 'Ubisoft' kelimesi 'Ubisoft Connect' olarak da geçebilir; en uzun eşleşmeyi dene
        for label in ([plat, 'Ubisoft Connect'] if plat == 'Ubisoft' else [plat, 'Game Pass'] if plat == 'Xbox' else [plat]):
            pat = re.compile(r'(?<![>\w])' + re.escape(label) + r'(?![\w<])')
            if pat.search(out):
                out = pat.sub(f'<a href="{url}" target="_blank" rel="noopener noreferrer" style="color:inherit;text-decoration:none;">{label}{SVG_EXT_LINK}</a>', out, count=1)
                break
    return out

def game_row(pp):
    m = GAME_RE.match(pp['text'])
    name, meta = m.group('name').strip(), m.group('meta').strip()
    links = [LINK_FIX.get(u, u) for u in pp['links']]
    href = links[0] if links else None
    c0 = render_game_cell(esc(name), studio_for(name), href)
    return [c0, render_genre_tags(*genres_for(name)), linkify_meta(esc(meta), links)], (name, meta, href)

def yt(vid, t=""):
    return (f'<div class="gp-yt-wrap" style="max-width:560px;margin:1.6em 0;"><iframe src="https://www.youtube.com/embed/{vid}" '
            f'title="{esc(t)}" frameborder="0" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share" '
            f'allowfullscreen loading="lazy" style="display:block;width:100%;aspect-ratio:16/9;height:auto;border:0;border-radius:12px;box-shadow:0 4px 14px rgba(0,0,0,0.5);"></iframe></div>')
def ytid(u):
    m = re.search(r'v=([A-Za-z0-9_-]{11})', u); return m.group(1) if m else ''

# ---- Enrichment metinleri (kural 16: klişe/hype yok, doğal dil, GFN varyasyonu) ----
editor_note = render_editor_note(
    "Monopoly: Star Wars Heroes vs. Villains, tanıdık Monopoly kurallarını Star Wars karakterleriyle birleştiriyor. "
    "Sıra tabanlı yapısı kısa oturumlara uygun; bulut üzerinden telefonda başladığın maçı televizyonda sürdürebilirsin.")
hatirlatma = render_highlight(
    "GeForce NOW oyun satmaz; sahip olduğun oyunları bulutta çalıştırır. Listedeki yapımları oynamak için "
    "ilgili mağazada (Steam, Epic Games Store, Xbox) geçerli bir lisansa ya da aboneliğe sahip olman gerekir.")
compact = render_compact_cta(
    "Monopoly: Star Wars Heroes vs. Villains",
    "Ayın ilk büyük eklemesi: sıra tabanlı Star Wars mücadelesi, kurulum beklemeden bulutta.",
    "GeForce NOW Paketleri", "https://gameplus.com.tr/gfn/paketler")
end_cta = render_end_cta(
    "Game+ ile bulutta oyun keyfine hazır mısın?",
    "Performance ve Ultimate paketleri kütüphanendeki GeForce NOW destekli yapımları donanım olmadan oynamanı sağlar. "
    "2.000'den fazla oyunu ve GFN Thursday'e eklenen yeni yapımları görmek için hemen kütüphaneye göz at!",
    btn2_label="GeForce NOW Oyunları", btn2_url="https://gameplus.com.tr/gfn/oyunlar")
TLDR_ITEMS = [
    "Temmuz boyunca GeForce NOW kütüphanesine 12 yeni oyun ekleniyor; açılışı Monopoly: Star Wars Heroes vs. Villains yapıyor.",
    "Bu hafta 2 oyun eklendi; Temmuz takviminin tamamı ve çıkış tarihleri yazıda.",
    "Haziran'dan katılan 10 yapım kütüphanede oynanmayı bekliyor.",
    "Listedeki oyunlara sahipsen, bulut üzerinden dilediğin cihazda devam edebilirsin.",
]
info = render_info_card([
    ("Temmuz'da eklenecek oyun", "12"),
    ("Bu hafta eklenen", "2"),
    ("Haziran'dan katılan", "10"),
    ("Geçen ay duyurulan", "18"),
])

# ---- Gövdeyi kur: yazar metni sırayla, oyun satırları tabloya ----
out, source_chunks, games_struct = [], [], []
i, n = 0, len(paras)
pending_rows = []
def flush_table():
    global pending_rows
    if pending_rows:
        out.append(render_table(["Oyun", "Tür", "Platform/Çıkış"], [r for r, _ in pending_rows]))
        games_struct.extend(g for _, g in pending_rows)
        pending_rows = []
prev_cards = []
while i < n:
    pp = paras[i]
    st, txt, html = pp['style'], pp['text'], pp['html']
    if st == 'Heading 1':
        flush_table(); out.append(f'<h1>{html}</h1>'); source_chunks.append(txt)
    elif st == 'Heading 2':
        flush_table()
        if 'Haziran’dan Kalan' in txt or "Haziran'dan Kalan" in txt:
            out.append(compact)          # Öne Çıkan (mock: Temmuz tablosu ile Haziran bölümü arasında)
        if 'Önceki Haftalarda' in txt:
            out.append(hatirlatma); out.append(end_cta)
        out.append(f'<h2>{html}</h2>'); source_chunks.append(txt)
        if 'Galakside' in txt:
            out.append(editor_note)
    elif re.match(r'^https?://(www\.)?youtube\.com/watch', txt):
        flush_table(); out.append(yt(ytid(txt)))
    elif txt.startswith('GFN Thursday') and pp['links']:
        m = re.search(r'\((\d+ \w+ \d{4})\)', txt)
        PREV_IMGS = {  # og:image (canlı blog kapakları)
            '25': 'https://gameplus.com.tr/gameplus-static-resources/images/blog/gfn-this-week-6-25-nv-blog.webp',
            '18': 'https://gameplus.com.tr/gameplus-static-resources/images/blog/gfn-this-week-6-18-blog-post-204.webp',
            '11': 'https://gameplus.com.tr/gameplus-static-resources/images/blog/gfn-this-week-6-11-blog.webp'}
        day = (m.group(1).split() if m else [''])[0]
        prev_cards.append({'url': pp['links'][0], 'date': m.group(1) if m else '', 'label': esc(txt),
                           'img': PREV_IMGS.get(day)})
    elif is_game_line(pp):
        row, g = game_row(pp); pending_rows.append((row, g))
    else:
        flush_table(); out.append(f'<p>{html}</p>'); source_chunks.append(txt)
    i += 1
flush_table()
if prev_cards: out.append(render_prev_weeks_cards(prev_cards))

body = "\n".join(out)
body, toc_items = inject_heading_ids(body)
toc = render_floating_toc([(l, t, a) for (l, t, a) in toc_items if l == 2])
tldr = render_tldr(TLDR_ITEMS, reading_time=estimate_reading_time(body))
body = body.replace('</h1>', '</h1>\n' + toc + tldr + info, 1)
body = ensure_leading_h1(body)
final = ANIMATED_BORDER_STYLE + "\n" + body

# ================= KONTROL NOKTALARI =================
print("=== 1) Çıktı yapısı ===")
ok1 = print_report(verify_output(final, blog_type="gfn", expect_faq=False))
print("=== 2) Yazar metni korundu mu (tablo satırları hariç, onlar yapısal dönüşüm) ===")
orig = "".join(f"<p>{esc(t)}</p>" for t in source_chunks)
ok2 = print_source_report(orig, final)
print("=== 3) Oyun satırları: isim + platform/tarih + link birebir mi ===")
_plain = re.sub(r'\s+', '', re.sub(r'<[^>]+>', ' ', final))   # tag'siz, boşluksuz metin (ikon tag'leri boşluk bırakır)
miss = 0
for (name, meta, href) in games_struct:
    okn = esc(name) in final
    okm = re.sub(r'\s+', '', meta) in _plain
    okh = (href is None) or (href in final)
    if not (okn and okm and okh):
        miss += 1; print(f"  ✗ {name}: isim={okn} meta={okm} link={okh}")
print(f"  {len(games_struct)-miss}/{len(games_struct)} oyun satırı birebir doğrulandı" + ("  ✓" if miss == 0 else ""))
print("=== 4) Embed + prev-weeks + H1 ===")
print(f"  embed: {final.count('youtube.com/embed')}/2 | prev kart: {len(prev_cards)}/3 | ilk başlık H1: {bool(re.match(chr(60)+'h1', body.lstrip()[:4]))}")
print("=== 5) GA4 id'leri ===")
for i_ in ['end-packages-button', 'end-games-button', 'featured-game-button']:
    print(f"  {i_}: {'✓' if i_ in final else '✗ YOK'}")

# ================= ÇIKTILAR =================
preview = embed_fonts(PAGE_HEAD.replace("__TITLE__", TITLE) + final + PAGE_FOOT)
open(os.path.join(OUT, f"ornek-blog-{SLUG}.html"), "w", encoding="utf-8").write(preview)
open(os.path.join(OUT, f"{SLUG}-body.html"), "w", encoding="utf-8").write(final)
doc = Document()
doc.add_heading("GFN Thursday 2 Temmuz 2026 - v10.1 Enriched HTML (CMS gövdesi)", level=1)
doc.add_paragraph("CMS gövdesine yapıştırılacak kod. Yazar metni birebir korunmuştur; yalnız enrichment eklenmiştir.")
for line in final.split("\n"):
    p = doc.add_paragraph(line)
    for r in p.runs: r.font.name = "Courier New"
doc.save(os.path.join(OUT, "GFN 2 Temmuz - v10 HTML.docx"))
print(f"\nÇıktılar: ornek-blog-{SLUG}.html | {SLUG}-body.html | GFN 2 Temmuz - v10 HTML.docx | chars: {len(final)}")

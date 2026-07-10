# -*- coding: utf-8 -*-
"""Battle Royale ham docx -> v10 enriched (metin KORUNUR, sadece EKLENİR). Checkpoint'li."""
import sys, os, re, zipfile
import xml.etree.ElementTree as ET
from docx import Document
sys.path.insert(0, "/Users/Erdo/.claude/skills/gameplus-blog-enrich/scripts")
from gameplus_blog_components import *

OUT = "/Users/Erdo/Desktop/Claude Projects/Dispatch"
DOCX = "/Users/Erdo/Downloads/En İyi Battle Royale Oyunları_ Son Kalan Olmaya Hazır mısın.docx"
TITLE = "En İyi Battle Royale Oyunları: Son Kalan Olmaya Hazır mısın?"
SLUG = "en-iyi-battle-royale-oyunlari-son-kalan-olmaya-hazir-misin"
NS = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
      'r':'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
def Wt(t): return '{%s}%s' % (NS['w'], t)
def Rt(t): return '{%s}%s' % (NS['r'], t)

z = zipfile.ZipFile(DOCX)
rels = {rel.get('Id'): rel.get('Target') for rel in ET.fromstring(z.read('word/_rels/document.xml.rels'))}
def esc(t): return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
def run_html(r):
    txt = ''.join(n.text or '' for n in r.iter(Wt('t')))
    if not txt: return ''
    rpr = r.find(Wt('rPr')); b = rpr is not None and rpr.find(Wt('b')) is not None
    i = rpr is not None and rpr.find(Wt('i')) is not None
    h = esc(txt)
    if b: h = '<strong>%s</strong>' % h
    if i: h = '<em>%s</em>' % h
    return h
def para_html(p):
    parts = []
    for ch in p._p:
        tag = ch.tag.split('}')[-1]
        if tag == 'r': parts.append(run_html(ch))
        elif tag == 'hyperlink':
            url = rels.get(ch.get(Rt('id')), '')
            inner = ''.join(run_html(r) for r in ch.findall(Wt('r')))
            parts.append('<a href="%s">%s</a>' % (url, inner) if url else inner)
    return ''.join(parts)

d = Document(DOCX)
blocks = []
for p in d.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    st = p.style.name
    kind = {'Heading 1':'h1','Heading 2':'h2','Heading 3':'h3'}.get(st, 'p')
    html = para_html(p).strip()
    if kind == 'p' and re.match(r'^https?://(www\.)?youtube\.com/watch', txt):
        kind, html = 'yt', txt
    blocks.append((kind, html, txt))

original_body = "\n".join(f'<{k if k in ("h1","h2","h3") else "p"}>{h}</{k if k in ("h1","h2","h3") else "p"}>'
                          for (k, h, t) in blocks if k != 'yt')

GENRE = {'Fortnite':'Aksiyon','PUBG: BATTLEGROUNDS':'FPS','Call of Duty: Warzone':'FPS',
         'Apex Legends':'FPS','Naraka: Bladepoint':'Dövüş','Garena Free Fire':'FPS','Battlefield REDSEC':'FPS'}
def ytid(u):
    m = re.search(r'v=([A-Za-z0-9_-]{11})', u); return m.group(1) if m else ''
def yt(vid, t=""):
    return (f'<div class="gp-yt-wrap" style="max-width:560px;margin:1.6em 0;"><iframe src="https://www.youtube.com/embed/{vid}" '
            f'title="{esc(t)}" frameborder="0" allow="accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share" '
            f'allowfullscreen loading="lazy" style="display:block;width:100%;aspect-ratio:16/9;height:auto;border:0;border-radius:12px;box-shadow:0 4px 14px rgba(0,0,0,0.5);"></iframe></div>')

n = len(blocks)
game_idx = [i for i in range(n) if blocks[i][0]=='h3' and i+1<n and blocks[i+1][0]=='yt']
game_names = [blocks[i][2] for i in game_idx]
first_game = game_idx[0] if game_idx else None
card_table = render_card_table("En İyi 7 Battle Royale Oyunu",
    [{'name':nm, 'badge':GENRE.get(nm,'FPS'), 'anchor':slugify(nm)} for nm in game_names])

cta_paketler = render_cta_paketler("Donanımın yetmese de son çembere kal",
    "Battle Royale yapımlarını yüksek performanslı bir bilgisayar almadan oynamak mümkün. GeForce NOW ile kütüphanendeki GeForce NOW destekli oyunları telefon, tablet ya da dizüstünden buluttan başlatabilirsin.")
end_cta = render_end_cta("Buluttan oynamaya hazır mısın?",
    "Performance ve Ultimate paketleriyle kütüphanendeki GeForce NOW destekli yapımları donanım yükseltmeden oyna.",
    btn2_label="GeForce NOW Oyunları", btn2_url="https://gameplus.com.tr/gfn/oyunlar", chip2="Oyunlar")

out = []
i = 0
while i < n:
    k, html, txt = blocks[i]
    if k == 'h1':
        out.append(f'<h1>{html}</h1>'); i += 1; continue
    if k == 'h2':
        if 'Sıkça Sorulan' in txt:
            pairs = []; j = i+1
            while j < n and blocks[j][0] != 'h2':
                if blocks[j][0]=='h3' and j+1<n and blocks[j+1][0]=='p':
                    pairs.append((blocks[j][2], blocks[j+1][1])); j += 2
                else: j += 1
            out.append(end_cta)
            out.append(f'<h2>{html}</h2>')
            out.append(render_faq_accordion(pairs))
            i = j; continue
        if 'Oyuncu Tarzına Göre' in txt:
            out.append(f'<h2>{html}</h2>'); i += 1
            if i < n and blocks[i][0]=='p': out.append(f'<p>{blocks[i][1]}</p>'); i += 1
            recs = []
            while i < n and blocks[i][0]=='p': recs.append(blocks[i][1]); i += 1
            out.append(render_list(recs, marker='dot')); continue
        if 'Son Çembere' in txt:      # 2. H2'den önce CTA Paketler
            out.append(cta_paketler)
        out.append(f'<h2>{html}</h2>'); i += 1; continue
    if k == 'h3' and i+1 < n and blocks[i+1][0]=='yt':
        if i == first_game: out.append(card_table)
        nm = txt
        out.append(render_game_h3_inline(slugify(nm), nm, GENRE.get(nm,'FPS'), None, "", level="h3"))
        out.append(yt(ytid(blocks[i+1][1]), nm))
        i += 2; continue
    if k == 'yt':
        out.append(yt(ytid(html))); i += 1; continue
    out.append(f'<p>{html}</p>'); i += 1

body = "\n".join(out)
body, toc_items = inject_heading_ids(body)
toc = render_floating_toc([(l,t,a) for (l,t,a) in toc_items])
tldr = render_tldr([
    "<strong>Battle Royale nedir:</strong> Onlarca oyuncunun ekipmansız başlayıp son kalan olmak için mücadele ettiği hayatta kalma türü.",
    "<strong>7 öne çıkan yapım:</strong> Fortnite, PUBG: BATTLEGROUNDS, Call of Duty: Warzone, Apex Legends, Naraka: Bladepoint, Garena Free Fire, Battlefield REDSEC.",
    "<strong>Tarzına göre seçim:</strong> Yapı kurma, taktiksel simülasyon, yüksek FPS, yakın dövüş veya mobil; her oyuncuya uygun bir seçenek var.",
    "<strong>Donanımın yetmiyorsa:</strong> GeForce NOW ile bu yapımların GeForce NOW destekli olanlarını buluttan oynayabilirsin.",
])
info = render_info_card([("İncelenen","7 Oyun"),("En Erişilebilir","Garena Free Fire"),
                         ("En Taktiksel","PUBG"),("Platform","PC · Mobil · Bulut")])
body = body.replace('</h1>', '</h1>\n' + toc + tldr + info, 1)
body = ensure_leading_h1(body)
final = ANIMATED_BORDER_STYLE + "\n" + body

# === KONTROL NOKTALARI ===
print("=== Çıktı yapısı kontrolü ===")
ok1 = print_report(verify_output(final, blog_type="general", n_games=len(game_names), expect_faq=True))
print("=== Kaynak metin korundu mu (halüsinasyon/silme) ===")
ok2 = print_source_report(original_body, final)
# link + embed korundu mu
src_links = set(rels.values())
kept = sum(1 for u in src_links if ('youtube' in u) or (u in final))
print(f"=== Link/embed: kaynaktaki {len(src_links)} URL'den {kept} tanesi çıktıda ===")
missing_links = [u for u in src_links if 'youtube' not in u and u not in final]
if missing_links: print("  EKSİK linkler:", missing_links)
print(f"=== embed sayısı: {final.count('youtube.com/embed')} (beklenen {len(game_names)}) ===")

# === ÇIKTILAR: preview + body-only + HTML-in-doc ===
from export_output import export
preview = embed_fonts(PAGE_HEAD.replace("__TITLE__", TITLE) + final + PAGE_FOOT)
export([{"title":TITLE,"slug":SLUG,"html":preview}], fmt="files-preview", out_dir=OUT)
open(os.path.join(OUT, f"{SLUG}-body.html"), "w", encoding="utf-8").write(final)
# HTML kodunu doc'a göm (manuel inceleme)
doc = Document()
doc.add_heading("Battle Royale - v10 Enriched HTML (CMS gövdesi)", level=1)
doc.add_paragraph("Aşağıdaki HTML, CMS gövdesine yapıştırılacak koddur. Yazar metni korunmuştur; yalnız enrichment eklenmiştir.")
for line in final.split("\n"):
    p = doc.add_paragraph(line);
    for r in p.runs: r.font.name = "Courier New"; r.font.size = None
    p.paragraph_format.space_after = None
doc.save(os.path.join(OUT, "Battle Royale - v10 HTML.docx"))
print("\nÇıktılar: ornek-blog-%s.html (önizleme) | %s-body.html | Battle Royale - v10 HTML.docx" % (SLUG, SLUG))
print("H1 ilk başlık mı:", final.lstrip().startswith('<style') and '<h1' in final.split('</style>')[1][:200])
